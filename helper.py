# Standard library imports
import os
import json
import re
import time
import subprocess
import tempfile
import shutil
from datetime import datetime
from typing import List, Dict, Any, Optional
from urllib.parse import urlparse, parse_qs

# Third-party imports
import requests
from bs4 import BeautifulSoup
from youtube_transcript_api import (
    YouTubeTranscriptApi,
    TranscriptsDisabled,
    NoTranscriptFound,
    CouldNotRetrieveTranscript
)
from pytube import YouTube, Playlist
import yt_dlp
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from google.generativeai import GenerativeModel
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import SystemMessage, HumanMessage
import cv2

# Configure Google API
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY', 'AIzaSyABCalHpFkvdfOu_K04MqKX3zAt-fndd1g')  # Fallback to hardcoded key if env var not set
os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY
genai.configure(api_key=GOOGLE_API_KEY)

# Constants
DEFAULT_CLIPS_DIR = "clips"
FFMPEG_CMD = r"C:\Users\my beast\AppData\Local\Microsoft\WinGet\Packages\Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe\ffmpeg-7.1.1-full_build\bin\ffmpeg.exe"
YTDLP_CMD = 'yt-dlp'

ydl_opts = {
        'quiet': True,  # Suppress unnecessary output
        'writeautomaticsub': True,  # Attempt to download automatic subtitles if available
        'subtitleslangs': ['en'],  # Specify the language for subtitles (e.g., English)
        'skip_download': True,  # Skip actual video download
        'extractaudio': False,
        'geo_bypass': True,
        'backoff_factor': 2,  # More aggressive backoff
    'max_sleep_interval': 10,# Don't download the audio
    'user_agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
    'nocheckcertificate': True,  # Skip SSL certificate verification
    'ignoreerrors': True,  # Continue on download errors
    'no_warnings': True,  # Suppress warnings
    'extract_flat': True,  # Don't extract full video info
    # Use web client instead of android to avoid token issues
    'extractor_args': {
        'youtube': {
            'player_client': ['web'],
            'player_skip': ['webpage', 'configs'],
        }
    }
}

"""#Transcription"""

def get_transcript(video_id: str) -> str:
    """
    Get transcript for a YouTube video using youtube-transcript-api.

    Args:
        video_id: YouTube video ID

    Returns:
        String containing the transcript text
    """
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        # Get transcript from YouTube directly
        transcript_data = YouTubeTranscriptApi.get_transcript(video_id)

        # Convert transcript data to a single string
        transcript_text = " ".join([entry['text'] for entry in transcript_data])

        print(f"✅ Retrieved transcript with {len(transcript_data)} entries")
        return transcript_text

    except Exception as e:
        print(f"❌ Failed to get transcript: {e}")
        return "No transcript available for this video."

def get_video_details(links):
    print("Processing links:", links)
    if isinstance(links, str):
        links = [links]
    
    video_details = []
    for link in links:
        try:
            print(f"Processing link: {link}")
            # Configure yt-dlp options
            ydl_opts = {
                'quiet': True,
                'no_warnings': True,
                'extract_flat': False,
                'force_generic_extractor': False
            }
            
            # Extract video information
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                try:
                    info = ydl.extract_info(link, download=False)
                    if not info:
                        print(f"Failed to extract info for {link}")
                        continue
                        
                    video_id = info.get('id')
                    print(f"Extracting info for video ID: {video_id}")
                    
                    # Get transcript
                    transcript = get_transcript(video_id)
                    
                    video_details.append({
                        'title': info.get('title', 'Unknown Title'),
                        'transcript': transcript,
                        'url': f"https://www.youtube.com/watch?v={video_id}"
                    })
                    print(f"Successfully processed video: {video_id}")
                    
                except Exception as e:
                    print(f"Error extracting video info: {str(e)}")
                    video_details.append({
                        'title': f'Error processing video {link}',
                        'transcript': 'Failed to retrieve video information. Please try again later.',
                        'url': link
                    })
                    
        except Exception as e:
            print(f"Error processing link {link}: {str(e)}")
            video_details.append({
                'title': f'Error processing video {link}',
                'transcript': 'Failed to retrieve video information. Please try again later.',
                'url': link
            })
    
    print(f"Completed processing {len(video_details)} videos")
    print("Retrieved video details:", video_details)
    return video_details

"""#Summarization"""

def summarize(videos):
    summaries = []
    summary_model = ChatGoogleGenerativeAI(
        model="gemini-2.0-flash-lite", 
        temperature=0.5,
        convert_system_message_to_human=True
    )

    for video in videos:
        transcript = video['transcript']
        messages = [
            SystemMessage(content="""
                You are an advanced AI model with deep expertise in the field of Literature and Science.
                Summarize the input YouTube transcript in no less than 500 words.
                Directly return the summary without referring to the transcript or the instructions.
            """),
            HumanMessage(content=transcript)
        ]
        try:
            summary = summary_model.invoke(messages).content
            summaries.append({
                "title": video['title'],
                "summary": summary
            })
            print(f"Successfully generated summary for: {video['title']}")
        except Exception as e:
            print(f"Error generating summary for {video['title']}: {str(e)}")
            summaries.append({
                "title": video['title'],
                "summary": f"Error generating summary: {str(e)}"
            })
        time.sleep(15)

    return summaries

"""#Translation"""

def translate(language, summaries):
    if language == "English":
        return summaries

    translation_model = ChatGoogleGenerativeAI(
        model="gemini-2.0-flash-lite", 
        temperature=1.0,
        convert_system_message_to_human=True
    )
    translated_summaries = []

    for summary in summaries:
        messages = [
            SystemMessage(
                content=f"""
                    You are a professional translator.
                    Translate the given English content into simple and general {language}.
                    Only return the translated content.
                """
            ),
            HumanMessage(content=summary['summary'])
        ]

        translated_text = translation_model.invoke(messages).content

        video = {
            "title": summary['title'],
            "summary": summary['summary'],
            "translated_summary": translated_text
        }
        translated_summaries.append(video)
        time.sleep(15)

    return translated_summaries

"""#Documentation"""

def markdown_to_html(markdown_text):
    """
    Convert markdown text to HTML.
    """
    return markdown.markdown(markdown_text)

def write_summaries_to_word(data, filename="summaries.docx"):
    """
    Generate a professionally formatted Word document containing video summaries.
    Supports multiple languages including RTL languages.
    
    Args:
        data: List of dictionaries containing video information
        filename: Output Word document filename
    """
    try:
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Video Summaries', 0)
        title.alignment = 1  # Center alignment
        
        # Add timestamp
        doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()  # Add spacing
        
        # Process each video
        for idx, item in enumerate(data, 1):
            title = item.get('title', f"Video {idx}")
            summary = item.get('summary', '')
            translated = item.get('translated_summary', '')
            video_id = item.get('videoId', '')
            language = item.get('language', 'English')
            
            # Add video title
            doc.add_heading(f"{idx}. {title}", level=1)
            
            # Add video ID if available
            if video_id:
                doc.add_paragraph(f"Video ID: {video_id}")
            
            # Add original summary
            if summary:
                doc.add_heading('Original Summary', level=2)
                doc.add_paragraph('English', style='Intense Quote')
                doc.add_paragraph(summary, style='Normal')
            
            # Add translated summary
            if translated:
                doc.add_heading('Translated Summary', level=2)
                doc.add_paragraph(language, style='Intense Quote')
                doc.add_paragraph(translated, style='Normal')
            
            # Add spacing between videos
            doc.add_paragraph()
        
        # Save the document
        doc.save(filename)
        print(f"✅ Word document saved as {filename}")
        return True

    except Exception as e:
        print(f"❌ Error generating Word document: {str(e)}")
        return False

"""#TL;DW"""

import os
import json
import subprocess
import tempfile
import shutil
import re
import concurrent.futures
from typing import List, Dict, Optional, Any

# Try imports that might not be available in all environments
try:
    from urllib.parse import urlparse, parse_qs
except ImportError:
    # Fallback for older Python versions
    from urllib.parse import urlparse
    from urllib.parse import parse_qs

# Constants to reduce repeated string creation
DEFAULT_CLIPS_DIR = "clips"
FFMPEG_CMD = 'ffmpeg'
YTDLP_CMD = 'yt-dlp'

def ask_gemini_for_key_clips(transcript: List[Dict[str, Any]]) -> Optional[List[Dict[str, Any]]]:
    """
    Uses LangChain and Gemini to extract key highlight timestamps from a transcript.
    Returns a list of dictionaries with start, end, and title for each clip.

    Args:
        transcript: List of dictionaries with 'start' and 'text' keys

    Returns:
        List of dictionaries with 'start', 'end', and 'title' for each key moment
    """
    if not transcript:
        print("❌ No transcript data provided")
        return None

    # Precompute the transcript string to avoid repetitive string operations
    # Use join instead of repeated concatenation for better performance
    formatted_lines = [f"[{int(entry['start'])}] {entry['text']}" for entry in transcript]
    formatted = "\n".join(formatted_lines)

    # Create prompt messages
    try:
        # Import these here to avoid startup overhead if function is not used
        from langchain_google_genai import ChatGoogleGenerativeAI
        from langchain_core.messages import SystemMessage, HumanMessage

        system_message = SystemMessage(content="""
        You are an intelligent video assistant.
        Given this transcript (with timestamps in seconds), identify EIGHT key moments that summarizes the video well.
        Make sure to include the most important parts of the video in the key moments.
        Make sure that the key moments are complete and self-contained and not cut off in the middle of a sentence or context.
        Each moment must include:
        * start: start time in seconds
        * end: end time in seconds (typically 30-60 seconds after start)
        * title: a short label (6–8 words)

        Return ONLY valid JSON in this format:
        [
          {"start": 60, "end": 90, "title": "Introduction to AI"},
          {"start": 150, "end": 180, "title": "Benefits of automation"}
        ]
        """)

        human_message = HumanMessage(content=f"Transcript:\n{formatted}")

        # Initialize Gemini model via LangChain with increased parallelism
        model = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash-lite",
            google_api_key=GOOGLE_API_KEY,
            temperature=0.7,
            convert_system_message_to_human=True
        )

        # Invoke the model with the prompt
        response = model.invoke([system_message, human_message])

        # Clean the response using more efficient string operations
        content = response.content
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0]
        elif "```" in content:
            content = content.split("```")[0]

        # Parse JSON response
        clips_data = json.loads(content.strip())

        print(f"✅ Identified {len(clips_data)} key moments in the video")
        for i, clip in enumerate(clips_data):
            print(f"  {i+1}. {clip['title']} ({clip['start']}-{clip['end']}s)")

        return clips_data

    except Exception as e:
        print(f"❌ Failed to identify key clips: {e}")
        return None

def download_clip(clip_info: Dict[str, Any], video_url: str, output_dir: str, index: int) -> Optional[str]:
    """
    Download a single clip from a YouTube video.

    Args:
        clip_info: Dictionary with 'start', 'end', and 'title'
        video_url: YouTube video URL
        output_dir: Directory to save clips
        index: Clip index for filename

    Returns:
        Path to downloaded clip file if successful, None otherwise
    """
    start_time = clip_info['start']
    end_time = clip_info['end']
    title = clip_info['title']

    # Create a safe filename from the title
    safe_title = "".join([c if c.isalnum() else "_" for c in title])
    output_file = os.path.join(output_dir, f"clip_{index:02d}_{safe_title}.mp4")

    # Format start and end times as HH:MM:SS
    def format_time(seconds):
        hrs = int(seconds // 3600)
        mins = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hrs:02d}:{mins:02d}:{secs:02d}"

    start_str = format_time(start_time)
    end_str = format_time(end_time)

    # Configure yt-dlp for optimized downloads
    download_cmd = [
        YTDLP_CMD,
        video_url,
        '--download-sections', f"*{start_str}-{end_str}",
        '--force-keyframes-at-cuts',
        '--output', output_file,
        '--quiet',
        '--no-warnings',
        '--no-check-certificate',
        '--no-playlist',
        '--no-playlist-reverse',
        '--extractor-args', 'youtube:player_client=web',
        '--format', 'mp4[height<=720]',  # Use a simple format that doesn't require ffmpeg
        '--ffmpeg-location', r"C:\Users\my beast\AppData\Local\Microsoft\WinGet\Packages\Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe\ffmpeg-7.1.1-full_build\bin"  # Specify FFmpeg location
    ]

    print(f"✂️ Extracting clip {index+1}: {title}")

    try:
        process = subprocess.run(download_cmd, check=True, timeout=180)
        if os.path.exists(output_file):
            return output_file
    except Exception as e:
        print(f"❌ Error downloading clip {index+1}: {str(e)}")

    return None

def download_clips(video_url: str, key_clips: List[Dict[str, Any]], output_dir: str = DEFAULT_CLIPS_DIR) -> List[str]:
    """
    Download specific clips from a YouTube video using yt-dlp's --download-sections.
    Uses parallel downloads for improved speed.

    Args:
        video_url: YouTube video URL
        key_clips: List of dictionaries with 'start', 'end', and 'title'
        output_dir: Directory to save clips

    Returns:
        List of paths to downloaded clip files
    """
    if not key_clips:
        print("❌ No key clips provided")
        return []

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Use ThreadPoolExecutor to download clips in parallel
    with concurrent.futures.ThreadPoolExecutor(max_workers=min(len(key_clips), os.cpu_count() * 2)) as executor:
        # Submit all download tasks
        future_to_index = {
            executor.submit(download_clip, clip, video_url, output_dir, i): i
            for i, clip in enumerate(key_clips)
        }

        # Collect results as they complete
    clip_files_by_index = {}

    # Collect results
    for future in concurrent.futures.as_completed(future_to_index):
        index = future_to_index[future]
        try:
            filepath = future.result()
            if filepath:
                clip_files_by_index[index] = filepath
        except Exception as e:
            print(f"❌ Error downloading clip {index+1}: {e}")

    # Order by index to ensure correct sequence
    clip_files = [clip for index, clip in sorted(clip_files_by_index.items())]

    print(f"✅ Extracted {len(clip_files)} clips")
    return clip_files

def split_text_into_lines(text, max_chars=30):
    """Split text into multiple lines with a maximum number of characters per line."""
    words = text.split()
    lines = []
    current_line = []
    current_length = 0
    
    for word in words:
        if current_length + len(word) + 1 <= max_chars:
            current_line.append(word)
            current_length += len(word) + 1
        else:
            lines.append(' '.join(current_line))
            current_line = [word]
            current_length = len(word)
    
    if current_line:
        lines.append(' '.join(current_line))
    
    return '\n'.join(lines)

def get_video_fps(video_path):
    """Get the frame rate of a video file using ffprobe."""
    try:
        # Get ffprobe path from ffmpeg path
        ffprobe_path = FFMPEG_CMD.replace('ffmpeg', 'ffprobe')
        
        # Run ffprobe to get frame rate
        cmd = [
            ffprobe_path,
            '-v', 'error',
            '-select_streams', 'v:0',
            '-show_entries', 'stream=r_frame_rate',
            '-of', 'default=noprint_wrappers=1:nokey=1',
            video_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        fps_str = result.stdout.strip()
        
        # Handle fractional frame rates (e.g., "30000/1001")
        if '/' in fps_str:
            num, den = map(int, fps_str.split('/'))
            return num / den
        return float(fps_str)
    except Exception as e:
        print(f"Warning: Could not detect frame rate for {video_path}, using default 30fps: {e}")
        return 30.0

def merge_clips(clips, output_path, font_path='static/TESLA.ttf'):
    """Merge video clips with transitions and titles using FFmpeg."""
    if not clips:
        return False

    # Create temp directory for processed clips
    temp_dir = tempfile.mkdtemp(prefix='merge_clips_')
    processed_clips = []
    
    try:
        # Process each clip to add title
        for i, clip in enumerate(clips, 1):
            clip_path = clip['path']
            clip_title = clip['title']
            
            # Sanitize title - remove or replace problematic characters
            clean_title = clip_title
            # Replace common problematic characters with spaces or remove them
            replacements = {
                "'": "",  # Remove apostrophes
                '"': "",  # Remove quotes
                ":": " -",  # Replace colons with dash
                ";": " -",  # Replace semicolons with dash
                "\\": "",  # Remove backslashes
                "/": " ",  # Replace forward slashes with space
                "|": " ",  # Replace pipe with space
                "<": "",  # Remove angle brackets
                ">": "",  # Remove angle brackets
                "*": "",  # Remove asterisks
                "?": "",  # Remove question marks
                "!": "",  # Remove exclamation marks
                "@": "at",  # Replace @ with 'at'
                "#": "",  # Remove hash
                "$": "",  # Remove dollar sign
                "%": "",  # Remove percent
                "^": "",  # Remove caret
                "&": "and",  # Replace & with 'and'
                "+": "plus",  # Replace + with 'plus'
                "=": "equals",  # Replace = with 'equals'
                "~": "",  # Remove tilde
                "`": "",  # Remove backtick
                "[": "",  # Remove square brackets
                "]": "",  # Remove square brackets
                "{": "",  # Remove curly braces
                "}": "",  # Remove curly braces
            }
            
            for char, replacement in replacements.items():
                clean_title = clean_title.replace(char, replacement)
            
            # Remove multiple spaces and trim
            clean_title = ' '.join(clean_title.split())
            
            # Split title into multiple lines
            wrapped_title = split_text_into_lines(clean_title)
            
            # Output path for processed clip
            processed_path = os.path.join(temp_dir, f'processed_{i:02d}.mp4')
            
            print(f"🎬 Adding title to clip {i}: {clean_title}")
            
            try:
                # Use FFmpeg to add title using drawtext filter
                add_title_cmd = [
            FFMPEG_CMD,
            '-hide_banner',
                    '-loglevel', 'error',
                    '-i', clip_path,
                    '-vf', f"drawtext=text='{wrapped_title}':fontfile='{font_path}':fontsize=(w/40):fontcolor=white:box=1:boxcolor=black@0.7:boxborderw=5:x=(w-text_w)/2:y=h*0.05",
                    '-c:v', 'libx264',
                    '-preset', 'medium',
                    '-crf', '23',
                    '-c:a', 'copy',  # Copy audio without re-encoding
                    processed_path
        ]

                subprocess.run(add_title_cmd, check=True)

                if os.path.exists(processed_path):
                    processed_clips.append(processed_path)
                else:
                    print(f"❌ Failed to create processed clip {i}")
                    return False

            except Exception as e:
                print(f"❌ Error processing clip {i}: {e}")
                return False

        if not processed_clips:
            print("❌ No valid clips to merge")
            return False

        # Create file list for FFmpeg
        file_list = os.path.join(temp_dir, 'file_list.txt')
        with open(file_list, 'w', encoding='utf-8') as f:
            for clip in processed_clips:
                f.write(f"file '{clip}'\n")

        # Use concat demuxer to merge clips
        merge_cmd = [
            FFMPEG_CMD,
            '-hide_banner',
            '-loglevel', 'error',
            '-f', 'concat',
            '-safe', '0',
            '-i', file_list,
            '-c:v', 'libx264',
            '-preset', 'medium',
            '-crf', '23',
            '-c:a', 'aac',
            '-b:a', '192k',
            output_path
        ]

        print("🎬 Merging clips...")
        subprocess.run(merge_cmd, check=True)
        print(f"✅ Successfully created highlight reel: {output_path}")
        return True

    except Exception as e:
        print(f"❌ Error creating highlight reel: {e}")
        return False
    finally:
        print("🧹 Cleaning up temporary files...")
        shutil.rmtree(temp_dir, ignore_errors=True)

def extract_video_id(video_url: str) -> Optional[str]:
    """
    Extract YouTube video ID from a URL.

    Args:
        video_url: YouTube video URL

    Returns:
        YouTube video ID
    """
    # Regular YouTube URL
    if "youtube.com/watch" in video_url:
        parsed_url = urlparse(video_url)
        video_id = parse_qs(parsed_url.query).get('v', [None])[0]
        return video_id

    # Shortened youtu.be URL
    elif "youtu.be/" in video_url:
        video_id = video_url.split("youtu.be/")[1].split("?")[0]
        return video_id

    # Embed URL
    elif "youtube.com/embed/" in video_url:
        video_id = video_url.split("youtube.com/embed/")[1].split("?")[0]
        return video_id

    # Already an ID (11-character string)
    elif re.match(r'^[A-Za-z0-9_-]{11}$', video_url):
        return video_url

    return None

def tldw(video_url: str, video_id: Optional[str] = None, output_file: Optional[str] = None,
         max_resolution: int = 720, parallel_downloads: bool = True) -> Optional[str]:
    """
    Main pipeline function to automatically generate a highlight reel from a YouTube video.
    Optimized for speed with parallel processing and efficient settings.

    Args:
        video_url: YouTube video URL
        video_id: YouTube video ID (optional, will be extracted from URL if not provided)
        output_file: Path to the output highlight reel (optional, will be generated if not provided)
        max_resolution: Maximum resolution for downloaded clips (default: 720p)
        parallel_downloads: Whether to download clips in parallel (default: True)

    Returns:
        Path to the highlight reel if successful, None otherwise
    """
    # Extract video ID if not provided
    if not video_id:
        video_id = extract_video_id(video_url)
        if not video_id:
            print(f"❌ Could not extract video ID from URL: {video_url}")
            return None

    print(f"🎬 Processing video: {video_id}")

    # Get transcript - this is a necessary sequential step
    transcript = get_transcript(video_id)
    if not transcript:
        print("❌ Failed to get transcript. Cannot continue.")
        return None

    # Identify key clips - this is a necessary sequential step
    key_clips = ask_gemini_for_key_clips(transcript)
    if not key_clips:
        print("❌ Failed to identify key clips. Cannot continue.")
        return None

    # Download clips - this can be parallelized
    clip_files = download_clips(video_url, key_clips)
    if not clip_files:
        print("❌ Failed to download clips. Cannot continue.")
        return None

    # Prepare clips data for merging
    clips_data = []
    for i, (clip_file, key_clip) in enumerate(zip(clip_files, key_clips)):
        clips_data.append({
            'path': clip_file,
            'title': key_clip['title']
        })

    # Merge clips into highlight reel
    if not output_file:
        output_file = f"highlight_{video_id}.mp4"
    result = merge_clips(clips_data, output_file)

    if result:
        print(f"🎉 Highlight reel created successfully: {output_file}")
        try:
            # Clean up individual clips
            for clip_file in clip_files:
                if os.path.exists(clip_file):
                    os.remove(clip_file)
            print("🧹 Cleaned up individual clips")
            
            # Return the output file path but don't delete it yet
            # It will be deleted after being sent to the user
            return output_file
        except Exception as e:
            print(f"⚠️ Warning: Error during cleanup: {e}")
            return output_file
    else:
        print("❌ Failed to create highlight reel")
        return None

# For Google Colab, don't use argparse directly since it conflicts with Colab's own arguments
# Instead, create a simple function to run with parameters

def run_tldw_colab(video_url, output_file=None, max_resolution=720, parallel_downloads=True):
    """
    Run the TLDW pipeline with the given parameters.
    This is a wrapper function for use in Google Colab.

    Args:
        video_url: YouTube video URL or ID
        output_file: Output file path (default: highlight_VIDEO_ID.mp4)
        max_resolution: Maximum resolution for downloaded clips (default: 720)
        parallel_downloads: Whether to download clips in parallel (default: True)

    Returns:
        Path to the highlight reel if successful, None otherwise
    """
    print(f"🚀 Starting TLDW for video: {video_url}")
    print(f"📊 Settings: max_resolution={max_resolution}, parallel_downloads={parallel_downloads}")

    return tldw(
        video_url=video_url,
        output_file=output_file,
        max_resolution=max_resolution,
        parallel_downloads=parallel_downloads
    )

"""#PipeLine"""

def pipeline(links, lang):
    print(f"Starting pipeline processing for {len(links)} links in {lang}")
    videos = []
    for link in links:
        print(f"Processing link: {link}")
        video_details = get_video_details(link)
        print(f"Retrieved video details: {video_details}")
        videos.extend(video_details)
    
    print(f"Total videos to process: {len(videos)}")
    final_content = []
    
    for i, video in enumerate(videos):
        print(f"Processing video {i+1}/{len(videos)}")
        try:
            print(f"Generating summary for: {video['title']}")
            # Create a list with a single video for summarize function
            video_list = [video]
            summary_result = summarize(video_list)
            if summary_result and len(summary_result) > 0:
                summary = summary_result[0]['summary']
                print("Summary generated successfully")
                
                if lang != 'English':
                    print(f"Translating summary to {lang}")
                    translated_summary = translate(lang, [{'title': video['title'], 'summary': summary}])
                    if translated_summary and len(translated_summary) > 0:
                        translated_summary = translated_summary[0]['translated_summary']
                    else:
                        translated_summary = None
                    print("Translation completed")
                else:
                    translated_summary = None
                    print("No translation needed for English")
                
                final_content.append({
                    'title': video['title'],
                    'summary': summary,
                    'translated_summary': translated_summary
                })
                print(f"Successfully processed video {i+1}")
            else:
                raise Exception("Failed to generate summary")
            
        except Exception as e:
            print(f"Error processing video {i+1}: {str(e)}")
            print(f"Error type: {type(e)}")
            import traceback
            print(f"Error traceback: {traceback.format_exc()}")
            final_content.append({
                'title': video['title'],
                'summary': f"Error processing video: {str(e)}",
                'translated_summary': None
            })
    
    print("Pipeline processing completed")
    return final_content

def get_gemini_model():
    """Get configured Gemini model."""
    try:
        model = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash-lite",
            temperature=0.5
        )
        return model
    except Exception as e:
        print(f"Error initializing Gemini model: {str(e)}")
        return None
